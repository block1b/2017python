# coding=utf8
# 改标签的字体，难点！！！
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.ns import qn
document = Document()
# paragraph = document.add_paragraph()
paragraph = document.add_heading(u'', 0)
run = paragraph.add_run(u"呵呵哒")

run.font.name = u"Adobe 楷体 Std R"
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'Adobe 楷体 Std R')

#保存文件
document.save('demo2.docx')
