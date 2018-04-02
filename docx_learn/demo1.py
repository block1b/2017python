# coding=utf
# 生成doc文件
from docx import Document
from docx.shared import Inches

# 打开空白文档
document = Document()
# 添加段落
paragraph = document.add_paragraph('123lor sit amet.')
paragraph2 = document.add_paragraph('456')
# paragraph2.style = 'ListNumber2'  # 会有警告
# 添加标题
document.add_heading('The REAL meaning of the universe')
# 添加表格
table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)  # 指定单元格
cell.text = 'parrot, possibly dead'  # 填充
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

# 添加图片
document.add_picture('image.png', width=Inches(1.0))
# style
paragraph3 = document.add_paragraph('Lorem ipsum ')
run = paragraph3.add_run('dolor')
run.bold = True  # 加粗
paragraph3.add_run(' sit amet.')

from docx import Document
# 居中图片
from docx.enum.text import WD_ALIGN_PARAGRAPH

my_image = document.add_picture('image.png')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加分页符
document.add_page_break()

# 保存文件
document.save('demo.docx')