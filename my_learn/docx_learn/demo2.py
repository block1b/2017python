# coding=utf8
# 实验报告doc模板

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


#打开文档
document = Document()

# 自定义heading1 style
styles = document.styles
new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = styles['Heading 1']
font = new_heading_style.font
font.name = 'Arial'
font.size = Pt(24)
document.add_paragraph('Header One', style='New Heading')


#加入不同等级的标题
document.add_heading('Document Title', 0)
document.add_heading(u'二级标题',1)
document.add_heading(u'二级标题',2)

#添加文本
paragraph = document.add_paragraph(u'添加了文本')
#设置字号
run = paragraph.add_run(u'设置字号')
run.font.size = Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'

#设置中文字体
run = paragraph.add_run(u'设置中文字体，')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    u'有序列表元素1',style='List Number'
)
document.add_paragraph(
    u'有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    u'无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
document.add_picture('image.png',width=Inches(1.25))

#增加分页
document.add_page_break()

#保存文件
document.save('demo.docx')
